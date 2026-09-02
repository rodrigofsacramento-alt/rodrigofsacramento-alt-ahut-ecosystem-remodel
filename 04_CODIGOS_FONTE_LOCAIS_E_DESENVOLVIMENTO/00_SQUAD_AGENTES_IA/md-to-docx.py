#!/usr/bin/env python3
"""Convert neurovendas lesson .md to .docx"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = sys.argv[1] if len(sys.argv) > 1 else '/opt/data/ahut-ecosystem/04_CODIGOS_FONTE_LOCAIS_E_DESENVOLVIMENTO/00_SQUAD_AGENTES_IA/NEUROVENDAS_AULA_3_HOTSEAT_UNIFICADO.md'
outpath = path.replace('.md', '.docx')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

with open(path) as f:
    for line in f:
        s = line.strip()
        if not s:
            continue
        if s.startswith('# ') and not s.startswith('## '):
            p = doc.add_heading(s[2:], 1); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif s.startswith('## '):
            doc.add_heading(s[3:], 2)
        elif s.startswith('### '):
            doc.add_heading(s[4:], 3)
        elif s.startswith('**') and s.endswith('**') and len(s) > 5:
            p = doc.add_paragraph(); run = p.add_run(s.strip('*')); run.bold = True; run.font.size = Pt(12)
        elif s in ('---', '——'):
            p = doc.add_paragraph(); run = p.add_run('─' * 50); run.font.size = Pt(8); run.font.color.rgb = RGBColor(150,150,150); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif s.startswith('*') and s.endswith('*') and len(s) > 3:
            p = doc.add_paragraph(); run = p.add_run(s.strip('*')); run.italic = True
        else:
            doc.add_paragraph(line.rstrip())

doc.save(outpath)
print(f'Saved: {outpath}')